from __future__ import annotations

import json

from app.api.file_ingest_service import (
    attach_file,
    build_attached_file_context_packet,
    get_file_context_summary,
    get_file_status,
)
from app.api.schemas.files import (
    FileKind,
    FileMemoryPosture,
    FileProcessingState,
)


def test_text_file_ingests_successfully_without_memory_promotion(tmp_path):
    source = tmp_path / "field-notes.txt"
    source.write_text(
        "Observation one.\nObservation two.\nObservation three.",
        encoding="utf-8",
    )

    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        conversation_id="conv_001",
        project_id="proj_001",
        ingest_root=ingest_root,
    )

    assert result.accepted is True
    assert result.blocked is False
    assert result.ready is True
    assert result.processing_state == FileProcessingState.READY
    assert result.file is not None
    assert result.file.file_kind == FileKind.TEXT
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.file.can_use_as_context is True
    assert result.file.can_promote_to_memory is False
    assert result.context_summary is not None
    assert result.context_summary.usable_as_context is True
    assert result.context_summary.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.context_summary.chunk_count >= 1

    raw_copy = ingest_root / "raw" / result.file_id / source.name
    chunks_path = ingest_root / "extracted" / result.file_id / "chunks.json"

    assert raw_copy.exists()
    assert chunks_path.exists()

    chunks_payload = json.loads(chunks_path.read_text(encoding="utf-8"))

    assert chunks_payload["file_id"] == result.file_id
    assert chunks_payload["display_name"] == source.name
    assert chunks_payload["chunk_count"] == result.context_summary.chunk_count
    assert chunks_payload["chunks"]


def test_markdown_file_ingests_successfully(tmp_path):
    source = tmp_path / "source.md"
    source.write_text(
        "# Heading\n\nThis is a markdown source file for local ingest.",
        encoding="utf-8",
    )

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is True
    assert result.ready is True
    assert result.blocked is False
    assert result.file is not None
    assert result.file.file_kind == FileKind.MARKDOWN
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.context_summary is not None
    assert result.context_summary.file_kind == FileKind.MARKDOWN
    assert result.context_summary.usable_as_context is True


def test_json_file_ingests_with_bounded_structure_summary(tmp_path):
    source = tmp_path / "data.json"
    source.write_text('{"site": "A", "values": [1, 2, 3]}', encoding="utf-8")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is True
    assert result.ready is True
    assert result.blocked is False
    assert result.file is not None
    assert result.file.file_kind == FileKind.JSON
    assert result.file.parser_used == "json_stdlib_parser"
    assert result.file.memory_promotion_allowed is False
    assert result.file.outward_sharing_allowed is False
    assert result.context_summary is not None
    assert result.context_summary.file_kind == FileKind.JSON
    assert result.context_summary.parser_used == "json_stdlib_parser"
    assert result.context_summary.retrieval_method == "bounded_selection"
    assert result.context_summary.memory_promotion_allowed is False
    assert result.context_summary.outward_sharing_allowed is False


def test_saved_html_ingests_without_fetching_or_executing_scripts(tmp_path):
    source = tmp_path / "saved.html"
    source.write_text(
        "<html><style>.x{}</style><script>steal()</script><body><h1>Visible</h1></body></html>",
        encoding="utf-8",
    )

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is True
    assert result.ready is True
    assert result.blocked is False
    assert result.file is not None
    assert result.file.file_kind == FileKind.HTML
    assert result.file.parser_used == "html_stdlib_text_parser"
    assert result.file.memory_promotion_allowed is False
    assert result.file.outward_sharing_allowed is False
    assert result.context_summary is not None
    assert result.context_summary.file_kind == FileKind.HTML
    assert result.context_summary.parser_used == "html_stdlib_text_parser"
    assert result.warnings


def test_unsupported_extension_is_blocked_without_claiming_ready_or_memory(tmp_path):
    source = tmp_path / "source.bin"
    source.write_bytes(b"pretend binary")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is False
    assert result.blocked is True
    assert result.ready is False
    assert result.processing_state == FileProcessingState.BLOCKED
    assert result.file is not None
    assert result.file.processing_state == FileProcessingState.BLOCKED
    assert result.file.file_kind == FileKind.UNSUPPORTED
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.file.can_use_as_context is False
    assert result.file.can_promote_to_memory is False
    assert result.errors
    assert "Unsupported file kind" in result.errors[0]


def test_pdf_without_parser_dependency_is_honestly_blocked(tmp_path):
    source = tmp_path / "source.pdf"
    source.write_bytes(b"%PDF-pretend")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    if result.ready:
        assert result.file is not None
        assert result.file.file_kind == FileKind.PDF
        return

    assert result.blocked is True
    assert result.file is not None
    assert result.file.file_kind == FileKind.PDF
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.errors
    assert (
        "pypdf or pdfplumber" in result.errors[0]
        or "could not be parsed locally" in result.errors[0]
        or "extractable text" in result.errors[0]
    )


def test_malformed_docx_is_honestly_blocked(tmp_path):
    source = tmp_path / "source.docx"
    source.write_bytes(b"pretend docx")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is False
    assert result.blocked is True
    assert result.ready is False
    assert result.file is not None
    assert result.file.file_kind == FileKind.DOCX
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.file.can_use_as_context is False
    assert result.errors
    assert "Traceback" not in result.errors[0]
    assert str(tmp_path) not in result.errors[0]


def test_valid_docx_ingests_when_python_docx_available(tmp_path):
    import importlib.util

    if importlib.util.find_spec("docx") is None:
        return

    from docx import Document

    source = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("Alpha DOCX paragraph.")
    document.add_paragraph("Beta local parser proof.")
    document.save(source)

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is True
    assert result.ready is True
    assert result.blocked is False
    assert result.file is not None
    assert result.file.file_kind == FileKind.DOCX
    assert result.file.parser_used == "docx_python_docx_text_parser"
    assert result.file.memory_promotion_allowed is False
    assert result.file.outward_sharing_allowed is False
    assert result.context_summary is not None
    assert result.context_summary.file_kind == FileKind.DOCX
    assert result.context_summary.parser_used == "docx_python_docx_text_parser"
    assert result.context_summary.chunk_count >= 1


def test_missing_file_fails_safely(tmp_path):
    source = tmp_path / "missing.txt"

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is False
    assert result.blocked is False
    assert result.ready is False
    assert result.processing_state == FileProcessingState.FAILED
    assert result.file is not None
    assert result.file.processing_state == FileProcessingState.FAILED
    assert result.errors
    assert "does not exist" in result.errors[0]


def test_directory_path_is_blocked_by_path_guard(tmp_path):
    source = tmp_path / "not-a-file"
    source.mkdir()

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is False
    assert result.blocked is True
    assert result.ready is False
    assert result.processing_state == FileProcessingState.BLOCKED
    assert result.errors
    assert "directory" in result.errors[0]


def test_oversized_supported_file_is_blocked(tmp_path):
    source = tmp_path / "large.txt"
    source.write_text("x" * 21, encoding="utf-8")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
        max_size_bytes=20,
    )

    assert result.accepted is False
    assert result.blocked is True
    assert result.ready is False
    assert result.processing_state == FileProcessingState.BLOCKED
    assert result.file is not None
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.errors
    assert "size limit" in result.errors[0]


def test_sensitive_env_like_path_is_blocked_before_ingest(tmp_path):
    source = tmp_path / ".env"
    source.write_text("TOKEN=secret", encoding="utf-8")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
    )

    assert result.accepted is False
    assert result.blocked is True
    assert result.ready is False
    assert result.file is not None
    assert result.file.can_use_as_context is False
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.file.blocked_reason
    assert "sensitive secret" in result.errors[0]
    assert str(tmp_path) not in result.errors[0]


def test_long_text_chunks_into_multiple_ordered_chunks(tmp_path):
    source = tmp_path / "long-notes.txt"
    source.write_text("abcdefghij" * 12, encoding="utf-8")

    result = attach_file(
        source,
        ingest_root=tmp_path / "ingest",
        chunk_char_limit=25,
    )

    assert result.accepted is True
    assert result.ready is True
    assert result.context_summary is not None
    assert result.context_summary.chunk_count > 1
    assert result.context_summary.selected_chunk_count == result.context_summary.chunk_count

    chunk_indexes = [
        chunk.chunk_index
        for chunk in result.context_summary.chunks
    ]

    assert chunk_indexes == list(range(result.context_summary.chunk_count))
    assert result.context_summary.memory_posture == FileMemoryPosture.NOT_MEMORY


def test_successful_ingest_writes_registry_result_json(tmp_path):
    source = tmp_path / "registry-notes.txt"
    source.write_text("Registry lookup should preserve ingest truth.", encoding="utf-8")
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    result_path = ingest_root / "extracted" / result.file_id / "ingest_result.json"

    assert result.ready is True
    assert result_path.exists()

    payload = json.loads(result_path.read_text(encoding="utf-8"))

    assert payload["file_id"] == result.file_id
    assert payload["ready"] is True
    assert payload["file"]["memory_posture"] == "not_memory"
    assert payload["context_summary"]["usable_as_context"] is True
    assert payload["context_summary"]["memory_posture"] == "not_memory"


def test_get_file_status_returns_persisted_ingest_result(tmp_path):
    source = tmp_path / "status-notes.txt"
    source.write_text("Status lookup should return the stored result.", encoding="utf-8")
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    status = get_file_status(
        result.file_id,
        ingest_root=ingest_root,
    )

    assert status is not None
    assert status.file_id == result.file_id
    assert status.ready is True
    assert status.file is not None
    assert status.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert status.file.can_use_as_context is True
    assert status.file.can_promote_to_memory is False
    assert status.context_summary is not None
    assert status.context_summary.memory_posture == FileMemoryPosture.NOT_MEMORY


def test_get_file_context_summary_returns_persisted_summary(tmp_path):
    source = tmp_path / "context-notes.md"
    source.write_text("# Context\n\nContext summary lookup should work.", encoding="utf-8")
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    summary = get_file_context_summary(
        result.file_id,
        ingest_root=ingest_root,
    )

    assert summary is not None
    assert summary.file_id == result.file_id
    assert summary.usable_as_context is True
    assert summary.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert summary.chunk_count >= 1
    assert summary.selected_chunk_count == summary.chunk_count


def test_missing_file_lookup_returns_none(tmp_path):
    ingest_root = tmp_path / "ingest"

    assert get_file_status("file_missing", ingest_root=ingest_root) is None
    assert get_file_context_summary("file_missing", ingest_root=ingest_root) is None


def test_build_attached_file_context_packet_returns_bounded_local_context(tmp_path):
    source = tmp_path / "bounded-context.md"
    source.write_text(
        "# Attached Context\n\n"
        "Alpha ecology note. Beta justice note. Gamma mapping note.\n\n"
        "Delta local context note. Epsilon stewardship note.",
        encoding="utf-8",
    )
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
        chunk_char_limit=35,
    )

    packet = build_attached_file_context_packet(
        [result.file_id, result.file_id],
        ingest_root=ingest_root,
        max_files=1,
        max_chunks_per_file=2,
        max_total_excerpt_chars=90,
    )

    assert packet["attached_files_are_memory"] is False
    assert packet["source"] == "user_selected_local_files"
    assert packet["locality"] == "local"
    assert packet["bounded"] is True
    assert packet["requested_file_ids"] == [result.file_id]
    assert packet["file_count"] == 1

    file_packet = packet["files"][0]

    assert file_packet["file_id"] == result.file_id
    assert file_packet["display_name"] == source.name
    assert file_packet["file_kind"] == "markdown"
    assert file_packet["memory_posture"] == "not_memory"
    assert file_packet["usable_as_context"] is True
    assert 1 <= file_packet["selected_chunk_count"] <= 2
    assert file_packet["chunks"]

    combined_excerpts = "\n".join(chunk["excerpt"] for chunk in file_packet["chunks"])

    assert "Alpha" in combined_excerpts
    assert len(combined_excerpts) <= 95


def test_build_attached_file_context_packet_reports_missing_files_honestly(tmp_path):
    packet = build_attached_file_context_packet(
        ["file_missing_001"],
        ingest_root=tmp_path / "ingest",
    )

    assert packet["attached_files_are_memory"] is False
    assert packet["file_count"] == 0
    assert packet["files"] == []
    assert packet["warnings"]
    assert "No local ingest record" in packet["warnings"][0]
    assert packet["errors"] == []




def test_csv_file_ingests_as_bounded_data_execution_input_without_text_chunks(tmp_path):
    source = tmp_path / "sites.csv"
    source.write_text(
        "site,temperature_c,ph\nA,12.5,7.1\nB,18.5,7.3\n",
        encoding="utf-8",
    )
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    assert result.accepted is True
    assert result.blocked is False
    assert result.ready is True
    assert result.processing_state == FileProcessingState.READY
    assert result.file is not None
    assert result.file.file_kind == FileKind.CSV
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.file.can_use_as_context is True
    assert result.file.can_promote_to_memory is False
    assert result.context_summary is not None
    assert result.context_summary.file_kind == FileKind.CSV
    assert result.context_summary.usable_as_context is True
    assert result.context_summary.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.context_summary.chunk_count == 0
    assert result.context_summary.selected_chunk_count == 0
    assert result.context_summary.chunks == []
    assert "data execution" in (result.context_summary.summary_note or "")

    raw_copy = ingest_root / "raw" / result.file_id / source.name
    chunks_path = ingest_root / "extracted" / result.file_id / "chunks.json"
    result_path = ingest_root / "extracted" / result.file_id / "ingest_result.json"

    assert raw_copy.exists()
    assert result_path.exists()
    assert not chunks_path.exists()


def test_build_attached_file_context_packet_routes_csv_to_data_files_not_text_chunks(tmp_path):
    source = tmp_path / "sites.csv"
    source.write_text(
        "site,temperature_c,ph\nA,12.5,7.1\nB,18.5,7.3\n",
        encoding="utf-8",
    )
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    packet = build_attached_file_context_packet(
        [result.file_id],
        ingest_root=ingest_root,
    )

    raw_copy = ingest_root / "raw" / result.file_id / source.name

    assert packet["attached_files_are_memory"] is False
    assert packet["source"] == "user_selected_local_files"
    assert packet["locality"] == "local"
    assert packet["bounded"] is True
    assert packet["requested_file_ids"] == [result.file_id]
    assert packet["used_file_ids"] == [result.file_id]
    assert packet["used_text_file_ids"] == []
    assert packet["used_data_file_ids"] == [result.file_id]
    assert packet["file_count"] == 1
    assert packet["text_file_count"] == 0
    assert packet["data_file_count"] == 1
    assert packet["files"] == []
    assert len(packet["data_files"]) == 1
    assert packet["warnings"] == []
    assert packet["errors"] == []

    data_file = packet["data_files"][0]

    assert data_file["file_id"] == result.file_id
    assert data_file["display_name"] == source.name
    assert data_file["file_name"] == source.name
    assert data_file["file_kind"] == "csv"
    assert data_file["source_kind"] == "attached_file"
    assert data_file["source_path"] == str(raw_copy)
    assert data_file["local_path"] == str(raw_copy)
    assert data_file["ready"] is True
    assert data_file["usable_as_context"] is True
    assert data_file["blocked"] is False
    assert data_file["memory_posture"] == "not_memory"
    assert data_file["chunk_count"] == 0
    assert data_file["selected_chunk_count"] == 0


def test_build_attached_file_context_packet_can_mix_text_context_and_csv_data_files(tmp_path):
    notes = tmp_path / "notes.md"
    notes.write_text("# Notes\n\nAlpha beta gamma.", encoding="utf-8")
    csv_source = tmp_path / "sites.csv"
    csv_source.write_text("site,value\nA,1\nB,2\n", encoding="utf-8")
    ingest_root = tmp_path / "ingest"

    notes_result = attach_file(notes, ingest_root=ingest_root)
    csv_result = attach_file(csv_source, ingest_root=ingest_root)

    packet = build_attached_file_context_packet(
        [notes_result.file_id, csv_result.file_id],
        ingest_root=ingest_root,
    )

    assert packet["file_count"] == 2
    assert packet["text_file_count"] == 1
    assert packet["data_file_count"] == 1
    assert packet["used_text_file_ids"] == [notes_result.file_id]
    assert packet["used_data_file_ids"] == [csv_result.file_id]
    assert len(packet["files"]) == 1
    assert len(packet["data_files"]) == 1
    assert packet["files"][0]["file_kind"] == "markdown"
    assert packet["files"][0]["chunks"]
    assert packet["data_files"][0]["file_kind"] == "csv"


def test_xlsx_file_ingests_as_bounded_data_execution_input_without_text_chunks(tmp_path):
    source = tmp_path / "workbook.xlsx"
    source.write_bytes(b"pretend xlsx content")
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    assert result.accepted is True
    assert result.blocked is False
    assert result.ready is True
    assert result.processing_state == FileProcessingState.READY
    assert result.file is not None
    assert result.file.file_kind == FileKind.XLSX
    assert result.file.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.file.can_use_as_context is True
    assert result.file.can_promote_to_memory is False
    assert result.context_summary is not None
    assert result.context_summary.file_kind == FileKind.XLSX
    assert result.context_summary.usable_as_context is True
    assert result.context_summary.memory_posture == FileMemoryPosture.NOT_MEMORY
    assert result.context_summary.chunk_count == 0
    assert result.context_summary.selected_chunk_count == 0
    assert result.context_summary.chunks == []
    assert "data execution" in (result.context_summary.summary_note or "")

    raw_copy = ingest_root / "raw" / result.file_id / source.name
    chunks_path = ingest_root / "extracted" / result.file_id / "chunks.json"
    result_path = ingest_root / "extracted" / result.file_id / "ingest_result.json"

    assert raw_copy.exists()
    assert result_path.exists()
    assert not chunks_path.exists()


def test_build_attached_file_context_packet_routes_xlsx_to_data_files_not_text_chunks(tmp_path):
    source = tmp_path / "workbook.xlsx"
    source.write_bytes(b"pretend xlsx content")
    ingest_root = tmp_path / "ingest"

    result = attach_file(
        source,
        ingest_root=ingest_root,
    )

    packet = build_attached_file_context_packet(
        [result.file_id],
        ingest_root=ingest_root,
    )

    raw_copy = ingest_root / "raw" / result.file_id / source.name

    assert packet["attached_files_are_memory"] is False
    assert packet["source"] == "user_selected_local_files"
    assert packet["locality"] == "local"
    assert packet["bounded"] is True
    assert packet["requested_file_ids"] == [result.file_id]
    assert packet["used_file_ids"] == [result.file_id]
    assert packet["used_text_file_ids"] == []
    assert packet["used_data_file_ids"] == [result.file_id]
    assert packet["file_count"] == 1
    assert packet["text_file_count"] == 0
    assert packet["data_file_count"] == 1
    assert packet["files"] == []
    assert len(packet["data_files"]) == 1
    assert packet["warnings"] == []
    assert packet["errors"] == []

    data_file = packet["data_files"][0]

    assert data_file["file_id"] == result.file_id
    assert data_file["display_name"] == source.name
    assert data_file["file_name"] == source.name
    assert data_file["file_kind"] == "xlsx"
    assert data_file["source_kind"] == "attached_file"
    assert data_file["source_path"] == str(raw_copy)
    assert data_file["local_path"] == str(raw_copy)
    assert data_file["ready"] is True
    assert data_file["usable_as_context"] is True
    assert data_file["blocked"] is False
    assert data_file["memory_posture"] == "not_memory"
    assert data_file["chunk_count"] == 0
    assert data_file["selected_chunk_count"] == 0
