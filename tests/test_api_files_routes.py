from __future__ import annotations

import app.api.file_ingest_service as file_ingest_service
from app.api.main import create_app
from tests.asgi_test_client import ASGITestClient


def _make_client() -> ASGITestClient:
    return ASGITestClient(create_app())


def _patch_ingest_root(monkeypatch, ingest_root):
    monkeypatch.setattr(
        file_ingest_service,
        "DEFAULT_INGEST_ROOT",
        ingest_root,
    )


def test_attach_text_file_returns_ok_envelope_without_memory_promotion(
    tmp_path,
    monkeypatch,
):
    source = tmp_path / "field-notes.txt"
    source.write_text(
        "Observation one.\nObservation two.\nObservation three.",
        encoding="utf-8",
    )
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
                "conversation_id": "conv_001",
                "project_id": "proj_001",
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert payload["result_type"] == "file_ingest_result"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["trace_summary"]["route_used"] == "files.attach"

    assert data["accepted"] is True
    assert data["blocked"] is False
    assert data["ready"] is True
    assert data["processing_state"] == "ready"
    assert data["file"]["file_kind"] == "text"
    assert data["file"]["memory_posture"] == "not_memory"
    assert data["file"]["can_use_as_context"] is True
    assert data["file"]["can_promote_to_memory"] is False
    assert data["context_summary"]["usable_as_context"] is True
    assert data["context_summary"]["memory_posture"] == "not_memory"


def test_attach_markdown_file_returns_ok_envelope(tmp_path, monkeypatch):
    source = tmp_path / "source.md"
    source.write_text(
        "# Field Notes\n\nMarkdown ingest should be local and bounded.",
        encoding="utf-8",
    )
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert payload["approval_state"] == "not_needed"
    assert data["ready"] is True
    assert data["file"]["file_kind"] == "markdown"
    assert data["context_summary"]["file_kind"] == "markdown"
    assert data["file"]["memory_posture"] == "not_memory"


def test_attach_unsupported_file_returns_blocked_envelope(tmp_path, monkeypatch):
    source = tmp_path / "source.bin"
    source.write_bytes(b"pretend binary")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "blocked"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "denied"
    assert payload["errors"]

    assert data["accepted"] is False
    assert data["blocked"] is True
    assert data["ready"] is False
    assert data["processing_state"] == "blocked"
    assert data["file"]["file_kind"] == "unsupported"
    assert data["file"]["memory_posture"] == "not_memory"
    assert data["file"]["can_use_as_context"] is False
    assert data["file"]["can_promote_to_memory"] is False


def test_attach_pdf_without_parser_dependency_returns_honest_blocked_envelope(
    tmp_path,
    monkeypatch,
):
    source = tmp_path / "source.pdf"
    source.write_bytes(b"%PDF-pretend")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    if payload["status"] == "ok":
        assert data["file"]["file_kind"] == "pdf"
        return

    assert payload["status"] == "blocked"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "denied"
    assert payload["errors"]
    assert (
        "pypdf or pdfplumber" in payload["errors"][0]
        or "could not be parsed locally" in payload["errors"][0]
        or "extractable text" in payload["errors"][0]
    )

    assert data["accepted"] is False
    assert data["blocked"] is True
    assert data["ready"] is False
    assert data["file"]["file_kind"] == "pdf"
    assert data["file"]["memory_posture"] == "not_memory"


def test_attach_malformed_docx_returns_blocked_not_route_error(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    source.write_bytes(b"pretend docx")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "blocked"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "denied"
    assert payload["errors"]
    assert "Traceback" not in payload["errors"][0]
    assert str(tmp_path) not in payload["errors"][0]

    assert data["accepted"] is False
    assert data["blocked"] is True
    assert data["ready"] is False
    assert data["file"]["file_kind"] == "docx"


def test_attach_valid_docx_returns_ok_envelope_when_parser_available(
    tmp_path,
    monkeypatch,
):
    import importlib.util

    if importlib.util.find_spec("docx") is None:
        return

    from docx import Document

    source = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("Alpha DOCX paragraph.")
    document.add_paragraph("Beta local parser proof.")
    document.save(source)
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert payload["approval_state"] == "not_needed"
    assert data["ready"] is True
    assert data["file"]["file_kind"] == "docx"
    assert data["file"]["parser_used"] == "docx_python_docx_text_parser"
    assert data["file"]["memory_promotion_allowed"] is False
    assert data["file"]["outward_sharing_allowed"] is False
    assert data["context_summary"]["parser_used"] == "docx_python_docx_text_parser"


def test_attach_json_file_returns_ok_envelope_with_local_parser(tmp_path, monkeypatch):
    source = tmp_path / "data.json"
    source.write_text('{"site": "A", "values": [1, 2, 3]}', encoding="utf-8")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert data["file"]["file_kind"] == "json"
    assert data["file"]["parser_used"] == "json_stdlib_parser"
    assert data["file"]["memory_promotion_allowed"] is False
    assert data["file"]["outward_sharing_allowed"] is False
    assert data["context_summary"]["parser_used"] == "json_stdlib_parser"
    assert data["context_summary"]["retrieval_method"] == "bounded_selection"


def test_attach_html_file_returns_ok_envelope_without_fetch(tmp_path, monkeypatch):
    source = tmp_path / "saved.html"
    source.write_text(
        "<html><script>steal()</script><body><h1>Visible</h1></body></html>",
        encoding="utf-8",
    )
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert data["file"]["file_kind"] == "html"
    assert data["file"]["parser_used"] == "html_stdlib_text_parser"
    assert data["file"]["memory_promotion_allowed"] is False
    assert data["file"]["outward_sharing_allowed"] is False
    assert data["context_summary"]["parser_used"] == "html_stdlib_text_parser"
    assert data["warnings"]


def test_attach_missing_file_returns_error_envelope(tmp_path, monkeypatch):
    source = tmp_path / "missing.txt"
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "error"
    assert payload["capability_state"] == "degraded"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["errors"]

    assert data["accepted"] is False
    assert data["blocked"] is False
    assert data["ready"] is False
    assert data["processing_state"] == "failed"
    assert data["file"]["processing_state"] == "failed"


def test_attach_request_body_must_be_json_object():
    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json=["not", "an", "object"],
        )

    assert response.status_code == 400

    payload = response.json()

    assert payload["status"] == "error"
    assert "must be a JSON object" in payload["errors"][0]
    assert payload["data"]["http_status_code"] == 400


def test_attach_source_path_is_required():
    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={},
        )

    assert response.status_code == 400

    payload = response.json()

    assert payload["status"] == "error"
    assert "source_path" in payload["errors"][0]
    assert payload["data"]["http_status_code"] == 400


def test_attach_positive_integer_fields_are_validated(tmp_path):
    source = tmp_path / "field-notes.txt"
    source.write_text("Small file.", encoding="utf-8")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
                "max_size_bytes": 0,
            },
        )

    assert response.status_code == 400

    payload = response.json()

    assert payload["status"] == "error"
    assert "max_size_bytes" in payload["errors"][0]
    assert payload["data"]["http_status_code"] == 400


def test_files_route_is_registered_in_root_app():
    with _make_client() as client:
        response = client.get("/")

    assert response.status_code == 200

    payload = response.json()

    assert "app.api.routes.files" in payload["data"]["registered_route_modules"]


def test_file_status_lookup_returns_saved_ingest_result(tmp_path, monkeypatch):
    source = tmp_path / "lookup-notes.txt"
    source.write_text("Lookup status should return stored ingest truth.", encoding="utf-8")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        attach_response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )
        file_id = attach_response.json()["data"]["file_id"]

        response = client.get(f"/files/{file_id}/status")

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert payload["result_type"] == "file_status"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["trace_summary"]["route_used"] == "files.status"

    assert data["file_id"] == file_id
    assert data["ready"] is True
    assert data["file"]["memory_posture"] == "not_memory"
    assert data["file"]["can_use_as_context"] is True
    assert data["file"]["can_promote_to_memory"] is False


def test_file_context_summary_lookup_returns_saved_summary(tmp_path, monkeypatch):
    source = tmp_path / "lookup-context.md"
    source.write_text("# Lookup\n\nContext summary should be retrievable.", encoding="utf-8")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        attach_response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
            },
        )
        file_id = attach_response.json()["data"]["file_id"]

        response = client.get(f"/files/{file_id}/context-summary")

    assert response.status_code == 200

    payload = response.json()
    data = payload["data"]

    assert payload["status"] == "ok"
    assert payload["result_type"] == "file_context_summary"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["trace_summary"]["route_used"] == "files.context_summary"

    assert data["file_id"] == file_id
    assert data["usable_as_context"] is True
    assert data["memory_posture"] == "not_memory"
    assert data["chunk_count"] >= 1
    assert data["selected_chunk_count"] == data["chunk_count"]


def test_unknown_file_status_lookup_returns_honest_error_envelope(monkeypatch, tmp_path):
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.get("/files/file_missing/status")

    assert response.status_code == 200

    payload = response.json()

    assert payload["status"] == "error"
    assert payload["result_type"] == "file_status"
    assert payload["capability_state"] == "degraded"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["errors"]
    assert payload["data"]["file_id"] == "file_missing"
    assert payload["data"]["found"] is False


def test_unknown_file_context_summary_lookup_returns_honest_error_envelope(
    monkeypatch,
    tmp_path,
):
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.get("/files/file_missing/context-summary")

    assert response.status_code == 200

    payload = response.json()

    assert payload["status"] == "error"
    assert payload["result_type"] == "file_context_summary"
    assert payload["capability_state"] == "degraded"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["errors"]
    assert payload["data"]["file_id"] == "file_missing"
    assert payload["data"]["found"] is False




def test_attach_csv_file_returns_ok_envelope_as_data_execution_input(
    tmp_path,
    monkeypatch,
):
    source = tmp_path / "sites.csv"
    source.write_text(
        "site,temperature_c,ph\n"
        "A,12.5,7.1\n"
        "B,18.5,7.3\n",
        encoding="utf-8",
    )
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
                "conversation_id": "conv_csv_001",
                "project_id": "proj_csv_001",
            },
        )

        assert response.status_code == 200

        payload = response.json()
        data = payload["data"]
        file_id = data["file_id"]

        context_response = client.get(f"/files/{file_id}/context-summary")

    assert payload["status"] == "ok"
    assert payload["result_type"] == "file_ingest_result"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["trace_summary"]["route_used"] == "files.attach"

    assert data["accepted"] is True
    assert data["blocked"] is False
    assert data["ready"] is True
    assert data["processing_state"] == "ready"
    assert data["file"]["file_kind"] == "csv"
    assert data["file"]["memory_posture"] == "not_memory"
    assert data["file"]["can_use_as_context"] is True
    assert data["file"]["can_promote_to_memory"] is False
    assert data["context_summary"]["file_kind"] == "csv"
    assert data["context_summary"]["usable_as_context"] is True
    assert data["context_summary"]["memory_posture"] == "not_memory"
    assert data["context_summary"]["chunk_count"] == 0
    assert data["context_summary"]["selected_chunk_count"] == 0
    assert data["context_summary"]["chunks"] == []
    assert "data execution" in data["context_summary"]["summary_note"]

    assert context_response.status_code == 200

    context_payload = context_response.json()
    context_data = context_payload["data"]

    assert context_payload["status"] == "ok"
    assert context_payload["result_type"] == "file_context_summary"
    assert context_payload["capability_state"] == "live"
    assert context_payload["locality"] == "local"
    assert context_payload["approval_state"] == "not_needed"
    assert context_payload["trace_summary"]["route_used"] == "files.context_summary"

    assert context_data["file_id"] == file_id
    assert context_data["file_kind"] == "csv"
    assert context_data["usable_as_context"] is True
    assert context_data["memory_posture"] == "not_memory"
    assert context_data["chunk_count"] == 0
    assert context_data["selected_chunk_count"] == 0
    assert context_data["chunks"] == []


def test_attach_xlsx_file_returns_ok_envelope_as_data_execution_input(
    tmp_path,
    monkeypatch,
):
    source = tmp_path / "workbook.xlsx"
    source.write_bytes(b"pretend xlsx content")
    _patch_ingest_root(monkeypatch, tmp_path / "ingest")

    with _make_client() as client:
        response = client.post(
            "/files/attach",
            json={
                "source_path": str(source),
                "conversation_id": "conv_xlsx_001",
                "project_id": "proj_xlsx_001",
            },
        )

        assert response.status_code == 200

        payload = response.json()
        data = payload["data"]
        file_id = data["file_id"]

        context_response = client.get(f"/files/{file_id}/context-summary")

    assert payload["status"] == "ok"
    assert payload["result_type"] == "file_ingest_result"
    assert payload["capability_state"] == "live"
    assert payload["locality"] == "local"
    assert payload["approval_state"] == "not_needed"
    assert payload["trace_summary"]["route_used"] == "files.attach"

    assert data["accepted"] is True
    assert data["blocked"] is False
    assert data["ready"] is True
    assert data["processing_state"] == "ready"
    assert data["file"]["file_kind"] == "xlsx"
    assert data["file"]["memory_posture"] == "not_memory"
    assert data["file"]["can_use_as_context"] is True
    assert data["file"]["can_promote_to_memory"] is False
    assert data["context_summary"]["file_kind"] == "xlsx"
    assert data["context_summary"]["usable_as_context"] is True
    assert data["context_summary"]["memory_posture"] == "not_memory"
    assert data["context_summary"]["chunk_count"] == 0
    assert data["context_summary"]["selected_chunk_count"] == 0
    assert data["context_summary"]["chunks"] == []
    assert "data execution" in data["context_summary"]["summary_note"]

    assert context_response.status_code == 200

    context_payload = context_response.json()
    context_data = context_payload["data"]

    assert context_payload["status"] == "ok"
    assert context_payload["result_type"] == "file_context_summary"
    assert context_payload["capability_state"] == "live"
    assert context_payload["locality"] == "local"
    assert context_payload["approval_state"] == "not_needed"
    assert context_payload["trace_summary"]["route_used"] == "files.context_summary"

    assert context_data["file_id"] == file_id
    assert context_data["file_kind"] == "xlsx"
    assert context_data["usable_as_context"] is True
    assert context_data["memory_posture"] == "not_memory"
    assert context_data["chunk_count"] == 0
    assert context_data["selected_chunk_count"] == 0
    assert context_data["chunks"] == []
