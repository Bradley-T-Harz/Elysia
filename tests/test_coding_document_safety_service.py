from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfWriter

from app.api.coding_document_adapter_service import extract_document_preview
from app.api.coding_document_safety_service import check_document_safety
from app.api.coding_document_type_registry import detect_document_type
from app.api.coding_path_guard_service import guard_workspace_path


def test_docx_container_is_allowed(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Hello document")
    document.save(path)

    result = check_document_safety(path, detect_document_type(path))

    assert result.allowed is True
    assert result.container_entry_count > 0
    assert result.byte_hash


def test_zip_slip_container_is_blocked(tmp_path: Path):
    path = tmp_path / "evil.docx"
    with ZipFile(path, "w") as archive:
        archive.writestr("../evil.txt", "nope")

    result = check_document_safety(path, detect_document_type(path))

    assert result.allowed is False
    assert result.blocked_reason == "zip_slip_path_traversal"


def test_macro_document_path_is_blocked_by_guard(tmp_path: Path):
    path = tmp_path / "evil.docm"
    path.write_bytes(b"not really a doc")

    guarded = guard_workspace_path(workspace_root=str(tmp_path), target_path=str(path))

    assert guarded.allowed is False
    assert guarded.reason == "unsupported_or_blocked_document_type"


def test_corrupted_container_is_blocked(tmp_path: Path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a zip container")

    result = check_document_safety(path, detect_document_type(path))

    assert result.allowed is False
    assert result.blocked_reason == "corrupted_or_invalid_container"


def test_encrypted_pdf_is_blocked(tmp_path: Path):
    path = tmp_path / "locked.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.encrypt("secret", algorithm="AES-256")
    with path.open("wb") as handle:
        writer.write(handle)

    result = check_document_safety(path, detect_document_type(path))

    assert result.allowed is False
    assert result.blocked_reason == "encrypted_document"


def test_oversized_document_is_blocked_with_small_limit(tmp_path: Path):
    path = tmp_path / "sample.pdf"
    path.write_bytes(b"%PDF-1.4\n" + b"x" * 1024 + b"\n%%EOF")

    result = check_document_safety(path, detect_document_type(path), max_document_bytes=10)

    assert result.allowed is False
    assert result.blocked_reason == "document_too_large"


def test_secret_like_text_is_redacted_from_extraction(tmp_path: Path):
    path = tmp_path / "secret.docx"
    document = Document()
    document.add_paragraph("API_KEY=super-secret-value")
    document.save(path)

    preview = extract_document_preview(path)

    assert preview.status == "completed"
    assert "super-secret-value" not in preview.text_preview
    assert preview.secret_scan_findings
    assert preview.redactions
