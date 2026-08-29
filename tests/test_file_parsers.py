from __future__ import annotations

import importlib.util

import app.api.file_text_extractors as file_text_extractors
from app.api.file_text_extractors import extract_file_text
from app.api.schemas.files import FileKind


def test_txt_and_markdown_parse_locally(tmp_path):
    txt = tmp_path / "notes.txt"
    txt.write_text("Alpha beta", encoding="utf-8")
    md = tmp_path / "notes.md"
    md.write_text("# Heading\n\nGamma delta", encoding="utf-8")

    txt_result = extract_file_text(txt, FileKind.TEXT)
    md_result = extract_file_text(md, FileKind.MARKDOWN)

    assert txt_result.ok is True
    assert "Alpha" in txt_result.text
    assert txt_result.parser_used == "plain_text_parser"
    assert md_result.ok is True
    assert "Heading" in md_result.text
    assert md_result.parser_used == "markdown_text_parser"


def test_json_parses_to_bounded_structure_summary(tmp_path):
    source = tmp_path / "data.json"
    source.write_text('{"site": "A", "values": [1, 2, 3]}', encoding="utf-8")

    result = extract_file_text(source, FileKind.JSON)

    assert result.ok is True
    assert result.parser_used == "json_stdlib_parser"
    assert "top_level_keys" in result.text
    assert "site" in result.text


def test_html_parser_strips_script_and_style_without_fetch(tmp_path):
    source = tmp_path / "saved.html"
    source.write_text(
        "<html><style>.x{}</style><script>steal()</script><body><h1>Visible</h1><p>Text</p></body></html>",
        encoding="utf-8",
    )

    result = extract_file_text(source, FileKind.HTML)

    assert result.ok is True
    assert result.parser_used == "html_stdlib_text_parser"
    assert "Visible" in result.text
    assert "Text" in result.text
    assert "steal" not in result.text
    assert result.warnings


def test_xlsx_parser_uses_openpyxl_when_available(tmp_path):
    if importlib.util.find_spec("openpyxl") is None:
        result = extract_file_text(tmp_path / "missing.xlsx", FileKind.XLSX)
        assert result.ok is False
        assert "openpyxl" in "; ".join(result.errors)
        return

    from openpyxl import Workbook

    source = tmp_path / "workbook.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sites"
    sheet.append(["site", "value"])
    sheet.append(["A", 1])
    workbook.save(source)

    result = extract_file_text(source, FileKind.XLSX)

    assert result.ok is True
    assert result.parser_used == "xlsx_openpyxl_summary_parser"
    assert "Sites" in result.text


def test_pdf_and_docx_are_honest_when_parser_dependency_missing(tmp_path, monkeypatch):
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-pretend")
    docx = tmp_path / "sample.docx"
    docx.write_bytes(b"pretend docx")

    monkeypatch.setattr(
        file_text_extractors.importlib.util,
        "find_spec",
        lambda name: None if name in {"pypdf", "pdfplumber", "docx"} else importlib.util.find_spec(name),
    )

    pdf_result = extract_file_text(pdf, FileKind.PDF)
    docx_result = extract_file_text(docx, FileKind.DOCX)

    assert pdf_result.ok is False
    assert "pypdf or pdfplumber" in "; ".join(pdf_result.errors)

    assert docx_result.ok is False
    assert "python-docx" in "; ".join(docx_result.errors)


def test_malformed_pdf_and_docx_return_safe_parser_errors(tmp_path):
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-pretend")
    docx = tmp_path / "sample.docx"
    docx.write_bytes(b"pretend docx")

    pdf_result = extract_file_text(pdf, FileKind.PDF)
    docx_result = extract_file_text(docx, FileKind.DOCX)

    assert pdf_result.ok is False
    assert pdf_result.parser_used.startswith("pdf_")
    assert pdf_result.errors
    assert "Traceback" not in "; ".join(pdf_result.errors)
    assert str(tmp_path) not in "; ".join(pdf_result.errors)

    assert docx_result.ok is False
    assert docx_result.parser_used.startswith("docx_")
    assert docx_result.errors
    assert "Traceback" not in "; ".join(docx_result.errors)
    assert str(tmp_path) not in "; ".join(docx_result.errors)


def test_valid_docx_extracts_text_when_python_docx_available(tmp_path):
    if importlib.util.find_spec("docx") is None:
        return

    from docx import Document

    source = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("Alpha DOCX paragraph.")
    document.add_paragraph("Beta local parser proof.")
    document.save(source)

    result = extract_file_text(source, FileKind.DOCX)

    assert result.ok is True
    assert result.parser_used == "docx_python_docx_text_parser"
    assert "Alpha DOCX paragraph." in result.text
    assert "Beta local parser proof." in result.text


def test_valid_blank_pdf_is_handled_without_crashing_when_pypdf_available(tmp_path):
    if importlib.util.find_spec("pypdf") is None:
        return

    from pypdf import PdfWriter

    source = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with source.open("wb") as handle:
        writer.write(handle)

    result = extract_file_text(source, FileKind.PDF)

    assert result.ok is False
    assert result.parser_used == "pdf_pypdf_text_parser"
    assert "extractable text" in "; ".join(result.errors)
