from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from app.api.main import create_app
from app.api.routes.coding_documents import (
    get_document_types,
    post_document_apply_approved,
    post_document_edit_plan,
    post_document_export_approved,
    post_document_export_plan,
    post_document_extract_preview,
    post_document_inspect,
)
from app.api.routes.coding_files import post_file_read_preview
from app.api.schemas.coding_documents import (
    CodingDocumentEditApplyRequest,
    CodingDocumentEditPlanRequest,
    CodingDocumentExportApplyRequest,
    CodingDocumentExportPlanRequest,
    CodingDocumentPathRequest,
)
from app.api.schemas.coding_files import CodingFileReadPreviewRequest
from tests.coding_approval_test_helpers import approval_fields_for_plan


async def _await_payload(coro):
    return await coro


def test_document_routes_registered_on_local_bridge():
    app = create_app()
    paths = set(app.openapi()["paths"])

    assert "/coding/document-types" in paths
    assert "/coding/document/inspect" in paths
    assert "/coding/document/extract-preview" in paths
    assert "/coding/document/export-plan" in paths
    assert "/coding/document/export-approved" in paths
    assert "/coding/document/edit-plan" in paths
    assert "/coding/document/apply-approved" in paths


def test_document_routes_inspect_extract_and_export_plan(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Route Document", level=1)
    document.add_paragraph("Route extraction works.")
    document.save(path)

    import asyncio

    types_payload = asyncio.run(_await_payload(get_document_types()))
    assert any(item["extension"] == ".docx" for item in types_payload["data"]["document_types"])

    payload = CodingDocumentPathRequest(workspace_root=str(tmp_path), file_path=str(path), approval_granted=True)
    inspect_payload = asyncio.run(_await_payload(post_document_inspect(payload)))
    assert inspect_payload["data"]["document"]["status"] == "completed"

    extract_payload = asyncio.run(_await_payload(post_document_extract_preview(payload)))
    assert "Route Document" in extract_payload["data"]["document"]["text_preview"]

    plan_payload = asyncio.run(
        _await_payload(
            post_document_export_plan(
                CodingDocumentExportPlanRequest(
                    workspace_root=str(tmp_path),
                    file_path=str(path),
                    approval_granted=True,
                    export_format="markdown",
                    target_path="sample.export.md",
                )
            )
        )
    )
    assert plan_payload["data"]["document_export_plan"]["status"] == "planned"
    export_plan = plan_payload["data"]["document_export_plan"]
    export_approval = approval_fields_for_plan(
        workspace_root=str(tmp_path),
        operation_kind="document_export",
        mutation_class="document_export",
        source_file=str(path),
        plan=SimpleNamespace(**export_plan),
    )

    export_result_payload = asyncio.run(
        _await_payload(
            post_document_export_approved(
                CodingDocumentExportApplyRequest(
                    workspace_root=str(tmp_path),
                    file_path=str(path),
                    approval_granted=True,
                    operator_approved=True,
                    export_format="markdown",
                    target_path="sample.export.md",
                    expected_source_hash=plan_payload["data"]["document_export_plan"]["source_hash"],
                    **export_approval,
                )
            )
        )
    )
    assert export_result_payload["data"]["document_export_result"]["status"] == "applied"

    edit_plan_payload = asyncio.run(
        _await_payload(
            post_document_edit_plan(
                CodingDocumentEditPlanRequest(
                    workspace_root=str(tmp_path),
                    file_path=str(path),
                    approval_granted=True,
                    operation="append_paragraph",
                    parameters={"text": "Added through route"},
                )
            )
        )
    )
    assert edit_plan_payload["data"]["document_edit_plan"]["status"] == "planned"
    edit_plan = edit_plan_payload["data"]["document_edit_plan"]
    edit_approval = approval_fields_for_plan(
        workspace_root=str(tmp_path),
        operation_kind="document_edit",
        mutation_class="document_edit",
        source_file=str(path),
        plan=SimpleNamespace(**edit_plan),
    )

    edit_result_payload = asyncio.run(
        _await_payload(
            post_document_apply_approved(
                CodingDocumentEditApplyRequest(
                    workspace_root=str(tmp_path),
                    file_path=str(path),
                    approval_granted=True,
                    operator_approved=True,
                    operation="append_paragraph",
                    parameters={"text": "Added through route"},
                    expected_source_hash=edit_plan_payload["data"]["document_edit_plan"]["source_hash"],
                    **edit_approval,
                )
            )
        )
    )
    assert edit_result_payload["data"]["document_edit_result"]["status"] == "applied"


def test_file_read_preview_routes_document_through_document_adapter(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Document via selected-file preview")
    document.save(path)

    import asyncio

    payload = asyncio.run(
        _await_payload(
            post_file_read_preview(
                CodingFileReadPreviewRequest(
                    workspace_root=str(tmp_path),
                    file_path=str(path),
                    approval_granted=True,
                )
            )
        )
    )

    preview = payload["data"]["file_preview"]
    assert preview["category"] == "document"
    assert preview["adapter"] == "document"
    assert "Document via selected-file preview" in preview["content_preview"]
