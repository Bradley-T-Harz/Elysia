"""DOCX adapter for bounded extraction and stable paragraph/table previews."""

from __future__ import annotations

from pathlib import Path
from typing import Any


def extract_docx_preview(path: Path, *, max_chars: int, max_tables: int, max_rows: int) -> dict[str, Any]:
    from docx import Document

    document = Document(str(path))
    props = document.core_properties
    metadata = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "title": props.title,
        "author": props.author,
        "subject": props.subject,
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
    }
    text_parts: list[str] = []
    outline: list[dict[str, Any]] = []
    provenance: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text or ""
        style_name = getattr(paragraph.style, "name", "") or ""
        if text and style_name.lower().startswith("heading"):
            outline.append({"kind": "heading", "paragraph": index, "style": style_name, "text": text[:300]})
        if text and len("\n".join(text_parts)) < max_chars:
            text_parts.append(f"[paragraph {index}] {text}")
            provenance.append({"kind": "paragraph", "paragraph": index, "chars": len(text)})

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables[:max_tables], start=1):
        rows: list[list[str]] = []
        for row in table.rows[:max_rows]:
            rows.append([cell.text for cell in row.cells])
        tables.append({"kind": "table", "table": table_index, "rows": rows, "row_count_previewed": len(rows)})
        provenance.append({"kind": "table", "table": table_index, "rows_previewed": len(rows)})

    return {
        "status": "completed",
        "metadata": metadata,
        "text_preview": "\n".join(text_parts)[:max_chars],
        "tables": tables,
        "outline": outline,
        "provenance": provenance,
        "warnings": [],
    }


__all__ = ("extract_docx_preview",)
