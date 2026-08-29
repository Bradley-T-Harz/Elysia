"""
Local text extractors for explicit user-selected files.

These extractors stay local-only. They do not execute JavaScript, fetch linked
resources, perform OCR, call cloud services, or promote content to memory.
"""

from __future__ import annotations

import csv
import importlib.util
import json
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from app.api.schemas.files import FileKind


@dataclass(frozen=True)
class ExtractedFileText:
    ok: bool
    text: str = ""
    parser_used: str = ""
    warnings: tuple[str, ...] = field(default_factory=tuple)
    errors: tuple[str, ...] = field(default_factory=tuple)
    metadata: dict[str, Any] = field(default_factory=dict)


class _SafeHTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._skip_depth = 0
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"} and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        text = " ".join(str(data or "").split())
        if text:
            self.parts.append(text)


def _read_text_with_fallback(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def _safe_parser_exception_message(path: Path, exc: Exception) -> str:
    """
    Keep parser errors compact and safe for route/service surfaces.
    """
    message = str(exc).replace(str(path), "selected file")
    return f"{type(exc).__name__}: {message}".strip()


def _extract_text(path: Path, file_kind: FileKind) -> ExtractedFileText:
    text = _read_text_with_fallback(path)
    parser = "markdown_text_parser" if file_kind == FileKind.MARKDOWN else "plain_text_parser"
    return ExtractedFileText(ok=True, text=text, parser_used=parser)


def _extract_json(path: Path) -> ExtractedFileText:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        return ExtractedFileText(
            ok=False,
            parser_used="json_stdlib_parser",
            errors=(f"Invalid JSON: {exc}",),
        )

    if isinstance(payload, dict):
        root_kind = "object"
        keys = list(payload.keys())[:20]
        summary = {
            "root_type": root_kind,
            "top_level_keys": keys,
            "top_level_key_count": len(payload),
        }
    elif isinstance(payload, list):
        root_kind = "array"
        summary = {
            "root_type": root_kind,
            "item_count": len(payload),
            "sample_item_types": sorted({type(item).__name__ for item in payload[:20]}),
        }
    else:
        summary = {"root_type": type(payload).__name__}

    text = json.dumps(summary, indent=2, sort_keys=True)
    return ExtractedFileText(
        ok=True,
        text=text,
        parser_used="json_stdlib_parser",
        metadata=summary,
    )


def _extract_html(path: Path) -> ExtractedFileText:
    parser = _SafeHTMLTextExtractor()
    parser.feed(_read_text_with_fallback(path))
    return ExtractedFileText(
        ok=True,
        text="\n".join(parser.parts),
        parser_used="html_stdlib_text_parser",
        warnings=("Scripts, styles, links, and external resources were not executed or fetched.",),
    )


def _extract_csv_summary(path: Path) -> ExtractedFileText:
    try:
        with path.open("r", encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            rows = []
            for index, row in enumerate(reader):
                rows.append(row)
                if index >= 25:
                    break
    except UnicodeDecodeError:
        with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
            reader = csv.reader(handle)
            rows = []
            for index, row in enumerate(reader):
                rows.append(row)
                if index >= 25:
                    break

    headers = rows[0] if rows else []
    summary = {
        "parser": "csv_stdlib_summary_parser",
        "columns": headers,
        "sample_row_count": max(0, len(rows) - 1),
    }
    return ExtractedFileText(
        ok=True,
        text=json.dumps(summary, indent=2, sort_keys=True),
        parser_used="csv_stdlib_summary_parser",
        metadata=summary,
    )


def _extract_xlsx_summary(path: Path) -> ExtractedFileText:
    if importlib.util.find_spec("openpyxl") is None:
        return ExtractedFileText(
            ok=False,
            parser_used="xlsx_openpyxl_parser_missing",
            errors=("XLSX parsing requires the local Python package: openpyxl.",),
        )

    from openpyxl import load_workbook  # type: ignore[import-not-found]

    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet_summaries: list[dict[str, Any]] = []
    for sheet in workbook.worksheets[:10]:
        sheet_summaries.append(
            {
                "sheet_name": sheet.title,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
            }
        )
    sheet_count = len(workbook.sheetnames)
    workbook.close()
    summary = {"sheets": sheet_summaries, "sheet_count": sheet_count}
    return ExtractedFileText(
        ok=True,
        text=json.dumps(summary, indent=2, sort_keys=True),
        parser_used="xlsx_openpyxl_summary_parser",
        metadata=summary,
    )


def _extract_pdf(path: Path) -> ExtractedFileText:
    """
    Extract text from a PDF using local parser dependencies only.

    This function is intentionally dependency-gated and local-only:
    - pypdf is preferred when available.
    - pdfplumber is an optional fallback when available.
    - OCR is not enabled.
    - cloud conversion is not used.
    - parser errors are returned as safe ExtractedFileText errors.
    """
    parser_errors: list[str] = []

    if importlib.util.find_spec("pypdf") is not None:
        from pypdf import PdfReader  # type: ignore[import-not-found]

        try:
            reader = PdfReader(str(path))
            if getattr(reader, "is_encrypted", False):
                return ExtractedFileText(
                    ok=False,
                    parser_used="pdf_pypdf_text_parser",
                    errors=("Encrypted PDF files are not supported by local ingest.",),
                )

            page_texts: list[str] = []
            pages = list(reader.pages)
            for index, page in enumerate(pages):
                extracted_text = page.extract_text() or ""
                if extracted_text.strip():
                    page_texts.append(f"Page {index + 1}:\n{extracted_text.strip()}")

            if not page_texts:
                return ExtractedFileText(
                    ok=False,
                    parser_used="pdf_pypdf_text_parser",
                    errors=("PDF did not contain extractable text. OCR is not enabled.",),
                    metadata={"page_count": len(pages)},
                )

            return ExtractedFileText(
                ok=True,
                text="\n\n".join(page_texts),
                parser_used="pdf_pypdf_text_parser",
                metadata={"page_count": len(pages)},
            )
        except Exception as exc:
            parser_errors.append(
                f"pypdf failed: {_safe_parser_exception_message(path, exc)}"
            )

    if importlib.util.find_spec("pdfplumber") is not None:
        import pdfplumber  # type: ignore[import-not-found]

        try:
            page_texts = []
            with pdfplumber.open(path) as pdf:
                for index, page in enumerate(pdf.pages):
                    extracted_text = page.extract_text() or ""
                    if extracted_text.strip():
                        page_texts.append(f"Page {index + 1}:\n{extracted_text.strip()}")
                page_count = len(pdf.pages)

            if not page_texts:
                return ExtractedFileText(
                    ok=False,
                    parser_used="pdf_pdfplumber_text_parser",
                    errors=("PDF did not contain extractable text. OCR is not enabled.",),
                    metadata={"page_count": page_count},
                )

            return ExtractedFileText(
                ok=True,
                text="\n\n".join(page_texts),
                parser_used="pdf_pdfplumber_text_parser",
                metadata={"page_count": page_count},
            )
        except Exception as exc:
            parser_errors.append(
                f"pdfplumber failed: {_safe_parser_exception_message(path, exc)}"
            )

    if parser_errors:
        return ExtractedFileText(
            ok=False,
            parser_used="pdf_local_text_parser_failed",
            errors=(f"PDF could not be parsed locally. {' | '.join(parser_errors)}",),
        )

    return ExtractedFileText(
        ok=False,
        parser_used="pdf_parser_missing",
        errors=("PDF text extraction requires one local Python package: pypdf or pdfplumber.",),
    )


def _extract_docx(path: Path) -> ExtractedFileText:
    if importlib.util.find_spec("docx") is None:
        return ExtractedFileText(
            ok=False,
            parser_used="docx_python_docx_parser_missing",
            errors=("DOCX text extraction requires the local Python package: python-docx.",),
        )

    from docx import Document  # type: ignore[import-not-found]

    try:
        document = Document(str(path))
    except Exception as exc:
        return ExtractedFileText(
            ok=False,
            parser_used="docx_python_docx_text_parser",
            errors=(f"DOCX could not be parsed locally: {_safe_parser_exception_message(path, exc)}",),
        )

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    if not parts:
        return ExtractedFileText(
            ok=False,
            parser_used="docx_python_docx_text_parser",
            errors=("DOCX did not contain extractable paragraph or table text.",),
        )
    return ExtractedFileText(
        ok=True,
        text="\n".join(parts),
        parser_used="docx_python_docx_text_parser",
        metadata={
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
    )


def extract_file_text(path: str | Path, file_kind: FileKind) -> ExtractedFileText:
    source = Path(path)
    if file_kind in {FileKind.TEXT, FileKind.MARKDOWN}:
        return _extract_text(source, file_kind)
    if file_kind == FileKind.JSON:
        return _extract_json(source)
    if file_kind == FileKind.HTML:
        return _extract_html(source)
    if file_kind == FileKind.CSV:
        return _extract_csv_summary(source)
    if file_kind == FileKind.XLSX:
        return _extract_xlsx_summary(source)
    if file_kind == FileKind.PDF:
        return _extract_pdf(source)
    if file_kind == FileKind.DOCX:
        return _extract_docx(source)
    return ExtractedFileText(
        ok=False,
        parser_used="unsupported_file_kind",
        errors=(f"Unsupported file kind for local text extraction: {file_kind.value}",),
    )


__all__ = ("ExtractedFileText", "extract_file_text")
